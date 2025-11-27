import pandas as pd
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import os
from openpyxl import load_workbook

TEMPLATE_PATH = "Diploma  nuevo 2025.docx"
EXCEL_PATH = "INFORMACIÓN DIPLOMAS.xlsx"
OUTPUT_FOLDER = "DIPLOMAS_GENERADOS"

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def replace_text_in_paragraph(paragraph, search, replace):
    """Reemplaza texto en un párrafo, incluso si está dividido en múltiples runs"""
    # Combinar todo el texto del párrafo
    full_text = ''.join(run.text for run in paragraph.runs)
    
    if search in full_text:
        # Limpiar todos los runs pero mantener el primero para preservar formato
        nuevo_texto = full_text.replace(search, replace)
        
        # Si hay runs, limpiar todos y poner el nuevo texto en el primero
        if paragraph.runs:
            for i, run in enumerate(paragraph.runs):
                if i == 0:
                    run.text = nuevo_texto
                else:
                    run.text = ''
        else:
            # Si no hay runs, crear uno nuevo
            paragraph.add_run(nuevo_texto)

def replace_text_in_cell(cell, search, replace):
    """Reemplaza texto en una celda, buscando en todos sus párrafos"""
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, search, replace)

def replace_text_in_xml_elements(element, search, replace):
    """Función recursiva para buscar y reemplazar en elementos XML"""
    from docx.oxml.ns import qn
    
    # Buscar en todos los elementos de texto (w:t)
    for t_elem in element.iter(qn('w:t')):
        if t_elem.text and search in t_elem.text:
            t_elem.text = t_elem.text.replace(search, replace)
    
    # Buscar en elementos de texto alternativos (pueden estar en diferentes namespaces)
    for t_elem in element.iter():
        if t_elem.tag.endswith('}t') and hasattr(t_elem, 'text') and t_elem.text:
            if search in t_elem.text:
                t_elem.text = t_elem.text.replace(search, replace)

def replace_text(doc, search, replace):
    """Reemplaza texto en todo el documento de forma exhaustiva"""
    # Reemplazar en párrafos del cuerpo del documento
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, search, replace)
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_cell(cell, search, replace)
    
    # Reemplazar en encabezados
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            for paragraph in header.paragraphs:
                replace_text_in_paragraph(paragraph, search, replace)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_text_in_cell(cell, search, replace)
        
        # Reemplazar en pies de página
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            for paragraph in footer.paragraphs:
                replace_text_in_paragraph(paragraph, search, replace)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_text_in_cell(cell, search, replace)
    
    # Reemplazar en TODOS los elementos XML del documento (cuadros de texto, formas, etc.)
    # Buscar recursivamente en todo el cuerpo del documento
    replace_text_in_xml_elements(doc.element.body, search, replace)

# Leer Excel usando pandas primero para obtener estructura
df = pd.read_excel(EXCEL_PATH, dtype=str, keep_default_na=False)

# Convertir columnas a MAYÚSCULAS para que coincidan con el Word
df.columns = df.columns.str.upper()

# Leer el Excel con openpyxl para obtener los valores formateados (con puntos)
print("Leyendo valores formateados del Excel...")
wb = load_workbook(EXCEL_PATH, data_only=False)
ws = wb.active

# Obtener nombres de columnas del Excel (primera fila)
excel_headers = []
for cell in ws[1]:
    header_val = str(cell.value).upper().strip() if cell.value else ''
    excel_headers.append(header_val)

print(f"Columnas encontradas en Excel: {excel_headers}")
print(f"Total de registros: {len(df)}")

# Si existe la columna N_DOCUMENTO, leer los valores formateados
if 'N_DOCUMENTO' in excel_headers:
    n_doc_col_idx = excel_headers.index('N_DOCUMENTO')
    
    # Función para formatear número con puntos como separador de miles
    def formatear_numero_con_puntos(numero):
        """Formatea un número con puntos como separadores de miles (formato colombiano)"""
        if numero is None:
            return ''
        
        # Convertir a string primero
        num_str = str(numero).strip()
        
        if num_str == '' or num_str.lower() == 'nan':
            return ''
        
        # Si ya tiene puntos, mantenerlo tal cual
        if '.' in num_str and not num_str.startswith('.'):
            # Verificar que los puntos estén en posiciones válidas (separadores de miles)
            # Si tiene puntos y parece un formato válido, mantenerlo
            partes = num_str.split('.')
            # Si todas las partes excepto la última tienen 3 dígitos, es formato válido
            es_formato_valido = all(len(p) == 3 for p in partes[:-1]) and len(partes[-1]) <= 3
            if es_formato_valido:
                return num_str
        
        try:
            # Si es numérico (int o float) o es un string de solo dígitos
            if isinstance(numero, (int, float)):
                num_int = int(numero)
            elif num_str.replace('.', '').isdigit():
                # Remover puntos existentes y convertir
                num_int = int(num_str.replace('.', ''))
            else:
                # Si no es numérico, devolver tal cual
                return num_str
            
            # Formatear agregando puntos cada 3 dígitos desde la derecha
            num_str_formateado = ''
            num_str_sin_puntos = str(num_int)
            
            # Agregar puntos de derecha a izquierda
            for i, digito in enumerate(reversed(num_str_sin_puntos)):
                if i > 0 and i % 3 == 0:
                    num_str_formateado = '.' + num_str_formateado
                num_str_formateado = digito + num_str_formateado
            
            return num_str_formateado
        except (ValueError, TypeError):
            # Si no se puede convertir, devolver el valor original
            return num_str
    
    # Leer cada valor de N_DOCUMENTO y formatearlo
    for idx in range(len(df)):
        cell = ws.cell(row=idx + 2, column=n_doc_col_idx + 1)  # +2 porque fila 1 es header, +1 porque openpyxl es base 1
        
        if cell.value is not None:
            # Obtener el valor de la celda
            valor_crudo = cell.value
            
            # Formatear el número con puntos
            valor_formateado = formatear_numero_con_puntos(valor_crudo)
            
            # Actualizar en el dataframe
            df.iloc[idx, df.columns.get_loc('N_DOCUMENTO')] = valor_formateado
            print(f"  Fila {idx + 1}: N_DOCUMENTO formateado como '{valor_formateado}'")

# Leer LUGAR_EXPEDICION directamente del Excel para preservar el formato exacto
lugar_expedicion_valores = {}
if 'LUGAR_EXPEDICION' in excel_headers:
    lugar_exp_col_idx = excel_headers.index('LUGAR_EXPEDICION')
    for idx in range(len(df)):
        cell = ws.cell(row=idx + 2, column=lugar_exp_col_idx + 1)  # +2 porque fila 1 es header, +1 porque openpyxl es base 1
        # Leer el valor exacto tal como está en el Excel (sin modificaciones, sin strip)
        if cell.value is not None:
            # Mantener el valor tal cual como está en el Excel, sin ninguna modificación
            # Usar str() directamente sin strip para preservar formato exacto
            valor_exacto = str(cell.value)
            lugar_expedicion_valores[idx] = valor_exacto
            print(f"  Fila {idx + 1}: LUGAR_EXPEDICION leído exactamente como: '{valor_exacto}'")
        else:
            lugar_expedicion_valores[idx] = ''

# Actualizar LUGAR_EXPEDICION en el DataFrame con valores exactos del Excel
if 'LUGAR_EXPEDICION' in df.columns and lugar_expedicion_valores:
    for idx in range(len(df)):
        if idx in lugar_expedicion_valores:
            df.iloc[idx, df.columns.get_loc('LUGAR_EXPEDICION')] = lugar_expedicion_valores[idx]

# Limpiar valores vacíos que pandas podría haber marcado como 'nan' string
# PERO no afectar LUGAR_EXPEDICION que ya tiene su valor preservado
for col in df.columns:
    if col != 'LUGAR_EXPEDICION':  # No modificar LUGAR_EXPEDICION
        df[col] = df[col].replace(['nan', 'NaN'], '')
    
wb.close()

for idx, (index, row) in enumerate(df.iterrows()):
    
    doc = Document(TEMPLATE_PATH)
    
    # Reemplazar cada campo
    for col in df.columns:
        placeholder = "{" + col + "}"   # Ej: {N_DOCUMENTO}
        
        # Manejo especial para LUGAR_EXPEDICION: mantener exactamente como está (sin modificaciones)
        if col == 'LUGAR_EXPEDICION':
            # Usar el índice de posición (idx) en lugar del índice del DataFrame para acceso directo
            valor = lugar_expedicion_valores.get(idx, '')
            # NO aplicar strip ni ninguna modificación al valor de LUGAR_EXPEDICION
        else:
            valor_raw = row[col]
            # Para otros campos, limpiar pero mantener el valor
            if pd.isna(valor_raw) or str(valor_raw).strip() == '' or str(valor_raw).strip().lower() == 'nan':
                valor = ""
            else:
                valor = str(valor_raw).strip()
        
        replace_text(doc, placeholder, valor)
        print(f"  Reemplazando {placeholder} con '{valor}'")

    # Nombre del archivo
    nombre_raw = row.get("NOMBRE_COMPLETO", f"SinNombre_{index+1}")
    nombre = str(nombre_raw).replace(" ", "_")

    output_path = os.path.join(OUTPUT_FOLDER, f"Diploma_{nombre}.docx")
    doc.save(output_path)
    print(f"✓ Diploma guardado: {output_path}")

print("✔ Diplomas generados correctamente en la carpeta DIPLOMAS_GENERADOS")
