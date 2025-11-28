from flask import Flask, render_template, request, send_file, jsonify
from flask_cors import CORS
import pandas as pd
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import os
from openpyxl import load_workbook
import zipfile
import tempfile
import shutil
from werkzeug.utils import secure_filename

app = Flask(__name__)
CORS(app)

# Obtener el directorio base donde está el script
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Configuración
UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'docx', 'xlsx', 'xls'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def replace_text_in_paragraph(paragraph, search, replace):
    """Reemplaza texto en un párrafo, incluso si está dividido en múltiples runs"""
    full_text = ''.join(run.text for run in paragraph.runs)
    
    if search in full_text:
        nuevo_texto = full_text.replace(search, replace)
        
        if paragraph.runs:
            for i, run in enumerate(paragraph.runs):
                if i == 0:
                    run.text = nuevo_texto
                else:
                    run.text = ''
        else:
            paragraph.add_run(nuevo_texto)

def replace_text_in_cell(cell, search, replace):
    """Reemplaza texto en una celda, buscando en todos sus párrafos"""
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, search, replace)

def replace_text_in_xml_elements(element, search, replace):
    """Función recursiva para buscar y reemplazar en elementos XML"""
    from docx.oxml.ns import qn
    
    for t_elem in element.iter(qn('w:t')):
        if t_elem.text and search in t_elem.text:
            t_elem.text = t_elem.text.replace(search, replace)
    
    for t_elem in element.iter():
        if t_elem.tag.endswith('}t') and hasattr(t_elem, 'text') and t_elem.text:
            if search in t_elem.text:
                t_elem.text = t_elem.text.replace(search, replace)

def replace_text(doc, search, replace):
    """Reemplaza texto en todo el documento de forma exhaustiva"""
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, search, replace)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_cell(cell, search, replace)
    
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            for paragraph in header.paragraphs:
                replace_text_in_paragraph(paragraph, search, replace)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_text_in_cell(cell, search, replace)
        
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            for paragraph in footer.paragraphs:
                replace_text_in_paragraph(paragraph, search, replace)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_text_in_cell(cell, search, replace)
    
    replace_text_in_xml_elements(doc.element.body, search, replace)

def formatear_numero_con_puntos(numero):
    """Formatea un número con puntos como separadores de miles (formato colombiano)"""
    if numero is None:
        return ''
    
    num_str = str(numero).strip()
    
    if num_str == '' or num_str.lower() == 'nan':
        return ''
    
    if '.' in num_str and not num_str.startswith('.'):
        partes = num_str.split('.')
        es_formato_valido = all(len(p) == 3 for p in partes[:-1]) and len(partes[-1]) <= 3
        if es_formato_valido:
            return num_str
    
    try:
        if isinstance(numero, (int, float)):
            num_int = int(numero)
        elif num_str.replace('.', '').isdigit():
            num_int = int(num_str.replace('.', ''))
        else:
            return num_str
        
        num_str_formateado = ''
        num_str_sin_puntos = str(num_int)
        
        for i, digito in enumerate(reversed(num_str_sin_puntos)):
            if i > 0 and i % 3 == 0:
                num_str_formateado = '.' + num_str_formateado
            num_str_formateado = digito + num_str_formateado
        
        return num_str_formateado
    except (ValueError, TypeError):
        return num_str

@app.route('/')
def index():
    return send_file('index.html')

@app.route('/app.js')
def app_js():
    return send_file('app.js', mimetype='application/javascript')

@app.route('/download-excel', methods=['GET'])
def download_excel():
    try:
        # Lista de posibles nombres del archivo Excel
        possible_names = [
            'INFORMACIÓN DIPLOMAS.xlsx',  # Nombre exacto con acento
            'INFORMACION DIPLOMAS.xlsx',  # Sin acento
            'informacion_diplomas.xlsx'   # Minúsculas con guiones bajos
        ]
        
        print(f"Directorio base: {BASE_DIR}")
        print(f"Directorio de trabajo actual: {os.getcwd()}")
        print(f"Buscando archivo Excel...")
        
        # Buscar el archivo en el directorio base
        for name in possible_names:
            # Intentar con ruta relativa primero
            if os.path.exists(name):
                full_path = os.path.abspath(name)
                print(f"Archivo encontrado (relativo): {full_path}")
                return send_file(
                    full_path,
                    mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    as_attachment=True,
                    download_name='INFORMACION_DIPLOMAS.xlsx'
                )
            
            # Intentar con ruta absoluta desde BASE_DIR
            full_path = os.path.join(BASE_DIR, name)
            if os.path.exists(full_path):
                print(f"Archivo encontrado (absoluto): {full_path}")
                return send_file(
                    full_path,
                    mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    as_attachment=True,
                    download_name='INFORMACION_DIPLOMAS.xlsx'
                )
            
            print(f"No encontrado: {name} (ruta: {full_path})")
        
        # Listar archivos en el directorio para debugging
        print(f"Archivos en directorio base:")
        try:
            for file in os.listdir(BASE_DIR):
                print(f"  - {file}")
        except Exception as e:
            print(f"Error al listar directorio: {e}")
        
        error_msg = 'Archivo Excel de ejemplo no encontrado. Verifica que el archivo "INFORMACIÓN DIPLOMAS.xlsx" esté en el directorio raíz del proyecto.'
        print(f"ERROR: {error_msg}")
        return jsonify({'error': error_msg}), 404
        
    except Exception as e:
        import traceback
        error_msg = f"Error al descargar Excel: {str(e)}"
        print(error_msg)
        print(traceback.format_exc())
        return jsonify({'error': error_msg}), 500

@app.route('/download-word', methods=['GET'])
def download_word():
    try:
        # Lista de posibles nombres del archivo Word
        possible_names = [
            'Diploma  nuevo 2025.docx',  # Nombre exacto con dos espacios
            'Diploma nuevo 2025.docx',   # Con un espacio
            'diploma nuevo 2025.docx',   # Minúsculas
            'diploma_nuevo_2025.docx'    # Con guiones bajos
        ]
        
        print(f"Directorio base: {BASE_DIR}")
        print(f"Directorio de trabajo actual: {os.getcwd()}")
        print(f"Buscando archivo Word...")
        
        # Buscar el archivo en el directorio base
        for name in possible_names:
            # Intentar con ruta relativa primero
            if os.path.exists(name):
                full_path = os.path.abspath(name)
                print(f"Archivo encontrado (relativo): {full_path}")
                return send_file(
                    full_path,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name='Diploma_nuevo_2025.docx'
                )
            
            # Intentar con ruta absoluta desde BASE_DIR
            full_path = os.path.join(BASE_DIR, name)
            if os.path.exists(full_path):
                print(f"Archivo encontrado (absoluto): {full_path}")
                return send_file(
                    full_path,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name='Diploma_nuevo_2025.docx'
                )
            
            print(f"No encontrado: {name} (ruta: {full_path})")
        
        # Listar archivos en el directorio para debugging
        print(f"Archivos en directorio base:")
        try:
            for file in os.listdir(BASE_DIR):
                print(f"  - {file}")
        except Exception as e:
            print(f"Error al listar directorio: {e}")
        
        error_msg = 'Archivo Word de ejemplo no encontrado. Verifica que el archivo "Diploma  nuevo 2025.docx" esté en el directorio raíz del proyecto.'
        print(f"ERROR: {error_msg}")
        return jsonify({'error': error_msg}), 404
        
    except Exception as e:
        import traceback
        error_msg = f"Error al descargar Word: {str(e)}"
        print(error_msg)
        print(traceback.format_exc())
        return jsonify({'error': error_msg}), 500

@app.route('/generate', methods=['POST'])
def generate_diplomas():
    try:
        # Verificar que se hayan enviado los archivos
        if 'template' not in request.files or 'excel' not in request.files:
            return jsonify({'error': 'Faltan archivos requeridos'}), 400
        
        template_file = request.files['template']
        excel_file = request.files['excel']
        
        if template_file.filename == '' or excel_file.filename == '':
            return jsonify({'error': 'No se seleccionaron archivos'}), 400
        
        if not allowed_file(template_file.filename) or not allowed_file(excel_file.filename):
            return jsonify({'error': 'Tipo de archivo no permitido'}), 400
        
        # Crear directorio temporal para trabajar
        temp_dir = tempfile.mkdtemp()
        output_folder = os.path.join(temp_dir, 'DIPLOMAS_GENERADOS')
        os.makedirs(output_folder, exist_ok=True)
        
        try:
            # Guardar archivos temporalmente
            template_path = os.path.join(temp_dir, secure_filename(template_file.filename))
            excel_path = os.path.join(temp_dir, secure_filename(excel_file.filename))
            
            template_file.save(template_path)
            excel_file.save(excel_path)
            
            # Leer Excel usando pandas
            df = pd.read_excel(excel_path, dtype=str, keep_default_na=False)
            df.columns = df.columns.str.upper()
            
            # Leer el Excel con openpyxl para obtener los valores formateados
            wb = load_workbook(excel_path, data_only=False)
            ws = wb.active
            
            # Obtener nombres de columnas del Excel
            excel_headers = []
            for cell in ws[1]:
                header_val = str(cell.value).upper().strip() if cell.value else ''
                excel_headers.append(header_val)
            
            # Formatear N_DOCUMENTO si existe
            if 'N_DOCUMENTO' in excel_headers:
                n_doc_col_idx = excel_headers.index('N_DOCUMENTO')
                
                for idx in range(len(df)):
                    cell = ws.cell(row=idx + 2, column=n_doc_col_idx + 1)
                    
                    if cell.value is not None:
                        valor_crudo = cell.value
                        valor_formateado = formatear_numero_con_puntos(valor_crudo)
                        df.iloc[idx, df.columns.get_loc('N_DOCUMENTO')] = valor_formateado
            
            # Leer LUGAR_EXPEDICION directamente del Excel
            lugar_expedicion_valores = {}
            if 'LUGAR_EXPEDICION' in excel_headers:
                lugar_exp_col_idx = excel_headers.index('LUGAR_EXPEDICION')
                for idx in range(len(df)):
                    cell = ws.cell(row=idx + 2, column=lugar_exp_col_idx + 1)
                    if cell.value is not None:
                        valor_exacto = str(cell.value)
                        lugar_expedicion_valores[idx] = valor_exacto
                    else:
                        lugar_expedicion_valores[idx] = ''
            
            # Actualizar LUGAR_EXPEDICION en el DataFrame
            if 'LUGAR_EXPEDICION' in df.columns and lugar_expedicion_valores:
                for idx in range(len(df)):
                    if idx in lugar_expedicion_valores:
                        df.iloc[idx, df.columns.get_loc('LUGAR_EXPEDICION')] = lugar_expedicion_valores[idx]
            
            # Limpiar valores vacíos
            for col in df.columns:
                if col != 'LUGAR_EXPEDICION':
                    df[col] = df[col].replace(['nan', 'NaN'], '')
            
            wb.close()
            
            # Generar diplomas
            for idx, (index, row) in enumerate(df.iterrows()):
                doc = Document(template_path)
                
                # Reemplazar cada campo
                for col in df.columns:
                    placeholder = "{" + col + "}"
                    
                    if col == 'LUGAR_EXPEDICION':
                        valor = lugar_expedicion_valores.get(idx, '')
                    else:
                        valor_raw = row[col]
                        if pd.isna(valor_raw) or str(valor_raw).strip() == '' or str(valor_raw).strip().lower() == 'nan':
                            valor = ""
                        else:
                            valor = str(valor_raw).strip()
                    
                    replace_text(doc, placeholder, valor)
                
                # Nombre del archivo
                nombre_raw = row.get("NOMBRE_COMPLETO", f"SinNombre_{index+1}")
                nombre = str(nombre_raw).replace(" ", "_")
                
                output_path = os.path.join(output_folder, f"Diploma_{nombre}.docx")
                doc.save(output_path)
            
            # Crear archivo ZIP
            zip_path = os.path.join(temp_dir, 'diplomas_generados.zip')
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(output_folder):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, output_folder)
                        zipf.write(file_path, arcname)
            
            # Enviar el archivo ZIP
            return send_file(
                zip_path,
                mimetype='application/zip',
                as_attachment=True,
                download_name='diplomas_generados.zip'
            )
            
        finally:
            # Limpiar archivos temporales
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    except Exception as e:
        import traceback
        error_details = str(e)
        print(f"Error en generate_diplomas: {error_details}")
        print(traceback.format_exc())
        return jsonify({'error': f'Error al procesar: {error_details}'}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug = os.environ.get('FLASK_ENV') == 'development'
    
    # Listar todas las rutas registradas para debug
    print("\n=== Rutas registradas ===")
    for rule in app.url_map.iter_rules():
        print(f"{rule.rule} -> {rule.endpoint} [{', '.join(rule.methods)}]")
    print("========================\n")
    
    app.run(host='0.0.0.0', port=port, debug=debug)

